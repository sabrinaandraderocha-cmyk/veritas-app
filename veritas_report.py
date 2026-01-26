from __future__ import annotations

import os
from typing import List, Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from reportlab.lib import colors

# Tenta importar Match, se não conseguir, define um dummy para não quebrar
try:
    from veritas_utils import Match
except ImportError:
    @dataclass
    class Match:
        query_chunk: str
        source_doc: str
        source_chunk: str
        score: float

# =========================================================
# UTILITÁRIOS GRÁFICOS (PDF)
# =========================================================

def _wrap_text(text: str, font_name: str, font_size: int, max_width: float):
    if not text: return []
    words = text.split()
    lines = []
    cur = ""
    for w in words:
        test = (cur + " " + w).strip()
        if stringWidth(test, font_name, font_size) <= max_width:
            cur = test
        else:
            if cur: lines.append(cur)
            cur = w
    if cur: lines.append(cur)
    return lines

def _new_page(c: canvas.Canvas, width: float, height: float, margin: float) -> float:
    c.showPage()
    return height - margin

def _draw_header(c: canvas.Canvas, title: str, margin: float, y: float) -> float:
    c.setFillColor(colors.darkblue)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(margin, y, title)
    c.setStrokeColor(colors.lightgrey)
    c.line(margin, y - 0.2*cm, A4[0] - margin, y - 0.2*cm)
    return y - 1.0 * cm

def _draw_kv(c: canvas.Canvas, margin: float, y: float, label: str, value: str) -> float:
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(colors.black)
    c.drawString(margin, y, f"{label}")
    
    c.setFont("Helvetica", 10)
    c.setFillColor(colors.darkgrey)
    c.drawString(margin + 4.5*cm, y, str(value))
    return y - 0.5 * cm

def _draw_section_title(c: canvas.Canvas, margin: float, y: float, txt: str) -> float:
    c.setFillColor(colors.black)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin, y, txt)
    return y - 0.6 * cm

# =========================================================
# 1. PDF: BIBLIOTECA (LOCAL)
# =========================================================
def generate_pdf_report(filepath, title, query_name, global_similarity, matches, params, disclaimer):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    margin = 2.0 * cm
    y = height - margin
    max_w = width - 2 * margin

    y = _draw_header(c, title, margin, y)
    y = _draw_kv(c, margin, y, "Arquivo:", query_name)
    y = _draw_kv(c, margin, y, "Similaridade Global:", f"{global_similarity*100:.1f}%")
    y -= 0.5 * cm

    # Parâmetros
    y = _draw_section_title(c, margin, y, "Parâmetros da Análise")
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.grey)
    for k, v in params.items():
        c.drawString(margin, y, f"• {k}: {v}")
        y -= 0.4 * cm
    
    y -= 0.5 * cm
    y = _draw_section_title(c, margin, y, "Correspondências Encontradas")
    
    if not matches:
        c.setFont("Helvetica", 10)
        c.setFillColor(colors.black)
        c.drawString(margin, y, "Nenhuma correspondência relevante encontrada.")
    else:
        for idx, m in enumerate(matches[:50], start=1):
            # Header do Match
            if y < margin + 4*cm: y = _new_page(c, width, height, margin)
            
            c.setFillColor(colors.black)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(margin, y, f"#{idx} | Fonte: {m.source_doc} ({m.score*100:.0f}%)")
            y -= 0.5 * cm
            
            # Textos
            c.setFont("Helvetica", 9)
            c.setFillColor(colors.darkgrey)
            
            # Texto Original
            c.drawString(margin, y, "Texto analisado:")
            y -= 0.4 * cm
            for line in _wrap_text(f'"{m.query_chunk}"', "Helvetica-Oblique", 9, max_w):
                c.drawString(margin + 0.5*cm, y, line)
                y -= 0.4 * cm
                if y < margin: y = _new_page(c, width, height, margin)

            # Texto Fonte
            y -= 0.2 * cm
            c.drawString(margin, y, "Texto fonte:")
            y -= 0.4 * cm
            for line in _wrap_text(f'"{m.source_chunk}"', "Helvetica", 9, max_w):
                c.drawString(margin + 0.5*cm, y, line)
                y -= 0.4 * cm
                if y < margin: y = _new_page(c, width, height, margin)
            
            y -= 0.5 * cm

    # Disclaimer no final
    if y < margin + 3*cm: y = _new_page(c, width, height, margin)
    y -= 1.0 * cm
    c.setStrokeColor(colors.lightgrey)
    c.line(margin, y, width-margin, y)
    y -= 0.5 * cm
    c.setFont("Helvetica", 8)
    c.setFillColor(colors.grey)
    for line in _wrap_text(disclaimer, "Helvetica", 8, max_w):
        c.drawString(margin, y, line)
        y -= 0.35 * cm

    c.save()

# =========================================================
# 2. PDF: INTERNET (WEB)
# =========================================================
def generate_web_pdf_report(filepath, title, query_name, profile, global_web_score, hits, disclaimer):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    margin = 2.0 * cm
    y = height - margin
    max_w = width - 2 * margin

    y = _draw_header(c, title, margin, y)
    y = _draw_kv(c, margin, y, "Arquivo:", query_name)
    y = _draw_kv(c, margin, y, "Score Web (Estimado):", f"{global_web_score*100:.1f}%")
    y -= 0.5 * cm

    y = _draw_section_title(c, margin, y, "Resultados da Busca")
    
    if not hits:
        c.drawString(margin, y, "Nenhuma ocorrência encontrada na web.")
    else:
        for idx, h in enumerate(hits[:30], start=1):
            if y < margin + 3*cm: y = _new_page(c, width, height, margin)
            
            # Título e Link
            c.setFillColor(colors.blue)
            c.setFont("Helvetica-Bold", 10)
            
            title_line = f"{idx}. {h.title[:60]}..." if len(h.title) > 60 else h.title
            c.drawString(margin, y, title_line)
            
            # Score
            c.setFillColor(colors.red)
            c.drawString(width - margin - 2*cm, y, f"{h.score*100:.0f}%")
            y -= 0.4 * cm
            
            # Link
            c.setFillColor(colors.grey)
            c.setFont("Helvetica", 8)
            c.drawString(margin, y, h.link[:90])
            y -= 0.4 * cm
            
            # Snippet
            c.setFillColor(colors.black)
            c.setFont("Helvetica", 9)
            for line in _wrap_text(h.snippet, "Helvetica", 9, max_w):
                c.drawString(margin, y, line)
                y -= 0.4 * cm
                if y < margin: y = _new_page(c, width, height, margin)
            
            y -= 0.3 * cm

    c.save()

# =========================================================
# 3. PDF: IA (HEURÍSTICO)
# =========================================================
def generate_ai_pdf_report(filepath, title, query_name, ai_result, disclaimer):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    margin = 2.0 * cm
    y = height - margin
    max_w = width - 2 * margin

    score = ai_result.get("score", 0)
    band = ai_result.get("band", ("Gray", "Desconhecido"))[1]

    y = _draw_header(c, title, margin, y)
    y = _draw_kv(c, margin, y, "Arquivo:", query_name)
    y = _draw_kv(c, margin, y, "Probabilidade IA:", f"{score:.0f}/100 ({band})")
    y -= 0.5 * cm

    # Métricas
    y = _draw_section_title(c, margin, y, "Indicadores Técnicos")
    c.setFont("Helvetica", 10)
    metrics = ai_result.get("metrics", {})
    
    metric_map = {
        "ttr": "Riqueza Vocabular (TTR)",
        "conn": "Densidade de Conectores",
        "vague": "Termos Vagos/Genéricos"
    }
    
    for k, v in metrics.items():
        label = metric_map.get(k, k)
        val = f"{v:.2f}"
        c.drawString(margin, y, f"• {label}: {val}")
        y -= 0.5 * cm

    y -= 0.5 * cm
    
    # Sentenças Sinalizadas
    y = _draw_section_title(c, margin, y, "Trechos com Padrões Artificiais")
    sentences = ai_result.get("flagged_sentences", [])
    
    c.setFont("Helvetica-Oblique", 9)
    if not sentences:
        c.drawString(margin, y, "Nenhuma frase atípica detectada.")
    else:
        for s in sentences[:15]:
            for line in _wrap_text(f"• {s}", "Helvetica-Oblique", 9, max_w):
                c.drawString(margin, y, line)
                y -= 0.4 * cm
                if y < margin: y = _new_page(c, width, height, margin)
            y -= 0.2 * cm

    # Disclaimer
    y -= 1.0 * cm
    if y < margin + 2*cm: y = _new_page(c, width, height, margin)
    
    c.setFont("Helvetica", 8)
    c.setFillColor(colors.grey)
    for line in _wrap_text(disclaimer, "Helvetica", 8, max_w):
        c.drawString(margin, y, line)
        y -= 0.35 * cm

    c.save()

# =========================================================
# 4. WORD: IA (HEURÍSTICO) - FALTAVA ESTA FUNÇÃO
# =========================================================
def generate_ai_docx_report(filepath, title, query_name, ai_result, disclaimer):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return # Se não tiver a lib, não faz nada

    doc = Document()
    
    # Título
    h = doc.add_heading(title, 0)
    
    # Meta
    p = doc.add_paragraph()
    p.add_run(f"Arquivo Analisado: ").bold = True
    p.add_run(query_name + "\n")
    
    score = ai_result.get("score", 0)
    band = ai_result.get("band", ("Gray", "Desconhecido"))[1]
    
    p.add_run(f"Score Heurístico: ").bold = True
    p.add_run(f"{score:.0f}/100 - {band}\n")

    # Métricas
    doc.add_heading("Métricas Técnicas", level=1)
    metrics = ai_result.get("metrics", {})
    for k, v in metrics.items():
        doc.add_paragraph(f"{k}: {v:.2f}", style='List Bullet')

    # Frases
    doc.add_heading("Trechos Sinalizados", level=1)
    sentences = ai_result.get("flagged_sentences", [])
    if not sentences:
        doc.add_paragraph("Nenhum padrão atípico encontrado.")
    else:
        for s in sentences:
            p = doc.add_paragraph(s)
            p.italic = True

    # Disclaimer
    doc.add_heading("Aviso Legal", level=2)
    doc.add_paragraph(disclaimer)

    doc.save(filepath)
