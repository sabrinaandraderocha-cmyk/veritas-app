from __future__ import annotations
import re
from dataclasses import dataclass
from typing import List, Tuple, Dict

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

def extract_text_from_txt_bytes(b: bytes) -> str:
    return b.decode("utf-8", errors="ignore")

def extract_text_from_docx_bytes(b: bytes) -> str:
    from docx import Document
    import io
    f = io.BytesIO(b)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                tx = cell.text.strip()
                if tx:
                    parts.append(tx)
    return "\n".join(parts)

def extract_text_from_pdf_bytes(b: bytes) -> str:
    from pypdf import PdfReader
    import io
    f = io.BytesIO(b)
    reader = PdfReader(f)
    parts = []
    for page in reader.pages:
        tx = page.extract_text() or ""
        tx = tx.strip()
        if tx:
            parts.append(tx)
    return "\n\n".join(parts)

def normalize_text(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[^\w\sáàâãéèêíïóôõöúçñ]", " ", s, flags=re.UNICODE)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def word_chunks(text: str, chunk_words: int = 60, stride_words: int = 20) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i+chunk_words]
        if len(chunk) < max(10, chunk_words//3):
            break
        chunks.append(" ".join(chunk))
        i += stride_words
    if not chunks and len(words) >= 10:
        chunks.append(" ".join(words))
    return chunks

@dataclass
class Match:
    query_chunk: str
    source_doc: str
    source_chunk: str
    score: float

def compute_matches(
    query_text: str,
    corpus_docs: Dict[str, str],
    chunk_words: int = 60,
    stride_words: int = 20,
    top_k_per_chunk: int = 1,
    threshold: float = 0.75,
) -> Tuple[float, List[Match]]:
    qn = normalize_text(query_text)
    q_chunks = word_chunks(qn, chunk_words=chunk_words, stride_words=stride_words)
    if not q_chunks:
        return 0.0, []

    doc_chunk_map: List[Tuple[str, str]] = []
    for docname, doctext in corpus_docs.items():
        dn = normalize_text(doctext)
        for ch in word_chunks(dn, chunk_words=chunk_words, stride_words=stride_words):
            doc_chunk_map.append((docname, ch))

    if not doc_chunk_map:
        return 0.0, []

    all_texts = q_chunks + [c for _, c in doc_chunk_map]
    vectorizer = TfidfVectorizer(ngram_range=(1,2), min_df=1)
    X = vectorizer.fit_transform(all_texts)
    Q = X[:len(q_chunks)]
    D = X[len(q_chunks):]

    sim = cosine_similarity(Q, D)
    best_scores = sim.max(axis=1)
    global_sim = float(best_scores.mean())

    matches: List[Match] = []
    for i in range(sim.shape[0]):
        row = sim[i]
        top_idx = row.argsort()[::-1][:top_k_per_chunk]
        for j in top_idx:
            score = float(row[j])
            if score >= threshold:
                docname, src_chunk = doc_chunk_map[int(j)]
                matches.append(Match(
                    query_chunk=q_chunks[i],
                    source_doc=docname,
                    source_chunk=src_chunk,
                    score=score
                ))

    uniq = {}
    for m in matches:
        key = (m.source_doc, m.source_chunk)
        if key not in uniq or m.score > uniq[key].score:
            uniq[key] = m
    matches = sorted(uniq.values(), key=lambda x: x.score, reverse=True)
    return global_sim, matches

def highlight_text(query_text: str, matches: List[Match]) -> str:
    out = query_text
    for m in matches[:20]:
        words = m.query_chunk.split()
        snippet = " ".join(words[: min(12, len(words))])
        if len(snippet) < 20:
            continue
        pattern = re.sub(r"\s+", r"\\s+", re.escape(snippet))
        try:
            out = re.sub(pattern, lambda mo: f"⟦{mo.group(0)}⟧", out, flags=re.IGNORECASE)
        except re.error:
            continue
    return out
